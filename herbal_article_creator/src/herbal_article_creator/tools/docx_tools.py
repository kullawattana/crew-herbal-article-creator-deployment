from __future__ import annotations
import os, re
from datetime import datetime
from typing import Optional, Dict, Any

# CrewAI BaseTool / Pydantic
from crewai.tools import BaseTool
from pydantic import BaseModel, Field

# Markdown -> HTML
import markdown as md
from bs4 import BeautifulSoup

# python-docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn           

class SaveMarkdownToDocxInput(BaseModel):
    markdown_text: str = Field(..., description="Markdown content to be converted to DOCX")
    output_file: Optional[str] = Field(
        default="research_paper.docx",
        description="The DOCX file path to save (e.g. outputs/research_paper.docx).",
    )
    research_topic: Optional[str] = Field(
        default="Herbal in Healthcare",
        description="Article topic for the cover page",
    )

class SaveMarkdownToDocxTool(BaseTool):
    """
    Tool: Convert Markdown to .docx file with cover, heading, list, table.
    """
    name: str = "save_markdown_to_docx"
    description: str = (
        "Convert Markdown to Word document (.docx) with formatting of headings, paragraphs,"
        "Bullet/numbered list and table with automatic cover generation"
    )
    args_schema: type[BaseModel] = SaveMarkdownToDocxInput

    @staticmethod
    def _ensure_paragraph_style(doc: Document, name: str, size_pt: int, bold: bool) -> str:
        """Create/restore paragraph style names in a document"""
        styles = doc.styles
        if name not in [s.name for s in styles]:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.font.size = Pt(size_pt)
            s.font.bold = bold
            s.font.name = "Times New Roman"
        else:
            s = styles[name]
            s.font.size = Pt(size_pt)
            s.font.bold = bold
            if not s.font.name:
                s.font.name = "Times New Roman"
        return name
    
    @staticmethod
    def _set_global_styles(doc: Document):
        """
        Style (Normal and Headings) 
        Font: Times New Roman, Font size and Font Color: Black for Normal, Dark Blue for Headings
        """
        # 1. Normal (General Body Text)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0) # Black
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # 2. Heading 1 to 6
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_font = h_style.font
                h_font.name = 'Times New Roman'
                h_font.bold = True
                h_font.color.rgb = RGBColor(54, 95, 145) # Dark Blue 
                h_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # size (Optional)
                if i == 1: h_font.size = Pt(16)
                elif i == 2: h_font.size = Pt(14)
                else: h_font.size = Pt(12)
    
    @staticmethod
    def _normalize_markdown(text: str) -> str:
        """
        Final Fix: Aggressively strip trailing noise (#, ##, > **) using Line-by-Line check.
        """
        # =========================================================
        # Start CLEANUP ZONE
        # =========================================================
        # 1. Line Endings (Fix \r)
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # 2. Line-by-Line
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            s = line.strip()
            
            # 2.1 delete (#, *, >) only lines
            if not s or re.fullmatch(r'[\s#*>]+', s):
                continue

            # 2.2 if there is content but trailing garbage (e.g., "ending. ##") -> remove tail
            # (Regex: remove space followed by # or * or > at the end)
            s = re.sub(r'[\s#*>]+$', '', s)
            
            cleaned_lines.append(s)
            
        text = '\n'.join(cleaned_lines)
        
        # 0.1) Protect fenced code blocks
        fences = []
        def _stash_fence(m):
            fences.append(m.group(0))
            return f"__FENCE_{len(fences)-1}__"
        text = re.sub(r"(?s)```.*?```", _stash_fence, text)

        # 0.2) Protect inline code
        inlines = []
        def _stash_inline(m):
            inlines.append(m.group(0))
            return f"__INLINE_{len(inlines)-1}__"
        text = re.sub(r"`[^`\n]+`", _stash_inline, text)

        # List topic headings to be formatted
        expected_headings = [
            "Herbal in Wellness Trends",
            "Scientific Deep-Dive",
            "Scientific Research",
            "Traditional Wisdom",
            "Safety, Regulatory, and Constraints",
            "Strategic Analysis and Product Opportunities",
            "Herbal Knowledge Summary",
            "Conclusion",
            "References",
            "Sources Consulted",
        ]

        # STEP 1: Pre-Split
        for h in expected_headings:
            text = re.sub(rf"([^\n])\s*({re.escape(h)})", r"\1\n\n\2", text)

        # STEP 2: Main Formatting
        for h in expected_headings:
            pattern = rf"(?m)^[\s#*]*{re.escape(h)}[\s*:]*\s+(?=\S)(.+?)[\s#*>]*$"
            text = re.sub(pattern, rf"### {h}\n\n\1", text)

        # Formatter Rules
        text = re.sub(r"(?m)(\S)[ \t]+(#{1,6}[ \t]+)", r"\1\n\n\2", text)
        text = re.sub(r"(?m)([^\n])\n(#{1,6}[ \t]+)", r"\1\n\n\2", text)
        text = re.sub(r"(?m)^(#{1,6}\s+.+?)$\n(?!\n)", r"\1\n\n", text)
        text = re.sub(
            r"(?m)^(#{1,6}\s+[\w\s\-]{3,100}?)\s{3,}([A-Z][\w\s]{50,})$",
            r"\1\n\n\2",
            text
        )
        text = re.sub(
            r"(?m)^(#{1,6}[ \t]+[^\n]*?)([ \t]+(?:[-*+]|\d+\.)[ \t]+)",
            r"\1\n\n\2",
            text,
        )
        text = re.sub(r"(?m)^>(\S)", r"> \1", text)
        text = re.sub(r"(?m)^(>\s*[^>\n]+:)\s*(>)\s*", r"\1\n>\n", text)
        text = re.sub(r"(?m)(\S)\n([ \t]*([-*+]|\d+\.)[ \t]+)", r"\1\n\n\2", text)
        text = re.sub(r"(?m)\s*(\[\d+\]\s*https?://)", r"\n\n\1", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        for i, block in enumerate(inlines):
            text = text.replace(f"__INLINE_{i}__", block)

        for i, block in enumerate(fences):
            text = text.replace(f"__FENCE_{i}__", block)
            
        # =========================================================
        # FINAL CLEANUP
        # =========================================================
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            s = line.strip()
            if not s or re.fullmatch(r'[\s#*>]+', s):
                continue
            s = re.sub(r'[\s#*>]+$', '', s)
            cleaned_lines.append(s)
        text = '\n'.join(cleaned_lines)
        # =========================================================
            
        return text

    def _run(self, markdown_text: str, output_file: str = "research_paper.docx",
             research_topic: str = "Machine Learning in Healthcare") -> Dict[str, Any]:
        """
        Convert markdown to .docx
        """
        try:
            outdir = os.path.dirname(output_file)
            if outdir:
                os.makedirs(outdir, exist_ok=True)
            
            # Normalize markdown
            markdown_text = self._normalize_markdown(markdown_text)

            # Markdown → HTML
            html = md.markdown(
                markdown_text,
                extensions=["extra", "tables", "toc", "sane_lists"],
            )
            soup = BeautifulSoup(html, "lxml")

            # Create document
            doc = Document()
            
            self._set_global_styles(doc)
            body_style = self._ensure_paragraph_style(doc, "CustomBody", size_pt=12, bold=False)

            # Cover page
            doc.add_heading(research_topic, level=0)
            doc.add_paragraph(
                f"Generated on {datetime.now().strftime('%B %d, %Y')}",
                style=body_style
            )
            doc.add_paragraph("Prepared by Academic Research by Crew with students of Thammasat University", style=body_style)
            doc.add_page_break()

            # Walk DOM and add content
            def handle_node(node):
                if node is None or isinstance(node, str):
                    return
                
                name = getattr(node, "name", None)
                if not name:
                    return

                if name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    level = int(name[1])
                    text = node.get_text(" ", strip=True)
                    doc.add_heading(text, level=level)                  
                    return

                elif name == "p":
                    text = node.get_text(strip=False)
                    if text.strip():
                        doc.add_paragraph(text, style='Normal')

                elif name in ["ul", "ol"]:
                    numbered = (name == "ol")
                    for li in node.find_all("li", recursive=False):
                        li_text = li.get_text(" ", strip=True)
                        self._add_list_item(doc, li_text, numbered=numbered, style_name=body_style)
                        for sub_list in li.find_all(["ul", "ol"], recursive=False):
                            for sub_li in sub_list.find_all("li", recursive=False):
                                self._add_list_item(
                                    doc,
                                    sub_li.get_text(" ", strip=True),
                                    numbered=(sub_list.name == "ol"),
                                    style_name=body_style,
                                )

                elif name == "table":
                    self._handle_table(node, doc)

                else:
                    # Process children for other tags
                    for child in node.children:
                        if getattr(child, "name", None):
                            handle_node(child)

            # Process content
            for el in soup.contents:
                if getattr(el, "name", None):
                    handle_node(el)

            # Save
            doc.save(output_file)
            return {
                "ok": True,
                "path": os.path.abspath(output_file),
                "message": f"Word document saved as {output_file}",
            }

        except Exception as e:
            return {"ok": False, "path": output_file, "message": f"Error: {e}"}